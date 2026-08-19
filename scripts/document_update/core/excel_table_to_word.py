"""
Excel 表格到 Word 表格整合模組
提供將 Excel 數據貼入已存在的 Word 表格的功能
"""

import os
import sys
import openpyxl
from docx import Document
from docx.shared import Pt


class ExcelTable2Word:
    """Excel 表格到 Word 表格整合類別"""
    
    def __init__(self):
        """初始化整合器"""
        pass
    
    def get_resource_path(self, relative_path):
        """獲取資源檔案的絕對路徑，適用於開發環境和 PyInstaller 打包環境"""
        if getattr(sys, 'frozen', False):
            # PyInstaller 打包環境 - 對於外部資源文件，要從 exe 同級目錄查找
            exe_dir = os.path.dirname(os.path.abspath(sys.argv[0]))
            # 檢查是否是 input 目錄下的文件
            if relative_path.startswith('input'):
                # input 目錄應該在 exe 的同級或上級目錄
                possible_paths = [
                    os.path.join(exe_dir, "..", relative_path),  # exe 在 dist 目錄中
                    os.path.join(exe_dir, relative_path),        # exe 和 input 在同一層
                ]
                for path in possible_paths:
                    abs_path = os.path.abspath(path)
                    if os.path.exists(abs_path):
                        return abs_path
                # 如果都找不到，返回默認路徑
                return os.path.join(exe_dir, relative_path)
            else:
                # 對於模板等內部資源，使用 _MEIPASS
                base_path = sys._MEIPASS
                return os.path.join(base_path, relative_path)
        else:
            # 開發環境
            base_path = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
            return os.path.join(base_path, relative_path)
    
    def get_output_path(self):
        """獲取輸出目錄的路徑，確保可寫入"""
        if getattr(sys, 'frozen', False):
            # PyInstaller 環境中，找到專案根目錄下的 output 目錄
            exe_dir = os.path.dirname(os.path.abspath(sys.argv[0]))
            # 嘗試找到專案根目錄（可能在上一層或同一層）
            possible_paths = [
                os.path.join(exe_dir, "..", "output"),  # exe 在 dist 目錄中，output 在上一層
                os.path.join(exe_dir, "output"),        # exe 和 output 在同一層
            ]
            
            for path in possible_paths:
                abs_path = os.path.abspath(path)
                if os.path.exists(abs_path):
                    return abs_path
            
            # 如果都找不到，就在 exe 同級目錄創建
            output_dir = os.path.join(exe_dir, "output")
        else:
            # 開發環境
            script_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
            output_dir = os.path.join(script_dir, "output")
        
        # 確保輸出目錄存在
        os.makedirs(output_dir, exist_ok=True)
        return output_dir
    
    def find_table_by_content(self, doc, search_text):
        """
        根據表格內容尋找目標表格
        
        Args:
            doc: Word 文件物件
            search_text: 要搜尋的文字
        
        Returns:
            表格索引，如果找不到則返回 -1
        """
        for table_idx, table in enumerate(doc.tables):
            for row in table.rows:
                for cell in row.cells:
                    if search_text in cell.text:
                        return table_idx
        return -1
    
    def paste_excel_to_word_table(self, excel_filename, sheet_name, word_filename, 
                                 table_identifier, cell_range=None, 
                                 start_position=(0, 0), clear_existing=False):
        """
        將 Excel 數據貼到 Word 文件中已存在的表格
        
        Args:
            excel_filename: Excel 檔案名稱（在 input 或 output 目錄中）
            sheet_name: Excel 工作表名稱
            word_filename: Word 檔案名稱（在 output 目錄中）
            table_identifier: 表格識別方式，可以是:
                            - int: 表格索引（從0開始）
                            - str: 要搜尋的文字內容
            cell_range: Excel 儲存格範圍，例如 "A1:E10"，None 表示整個有效範圍
            start_position: 在 Word 表格中開始貼上的位置 (row, col)
            clear_existing: 是否清除目標表格的現有內容
        
        Returns:
            bool: 操作是否成功
        """
        
        # 儲存當前處理的工作表名稱，用於判斷是否需要位置合併
        self.current_sheet_name = sheet_name
        
        try:
            # 尋找 Excel 檔案路徑 - 統一使用與 main.py 和 report_generator.py 相同的邏輯
            excel_path = None
            
            if getattr(sys, 'frozen', False):
                # PyInstaller 環境：Excel 檔案在 exe 同級或上級的 input 目錄
                exe_dir = os.path.dirname(os.path.abspath(sys.argv[0]))
                possible_paths = [
                    os.path.join(exe_dir, "..", "input", excel_filename),  # exe 在 dist 目錄中
                    os.path.join(exe_dir, "input", excel_filename),        # exe 和 input 在同一層
                ]
                
                for path in possible_paths:
                    abs_path = os.path.abspath(path)
                    if os.path.exists(abs_path):
                        excel_path = abs_path
                        break
            else:
                # 開發環境：直接使用相對路徑
                script_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
                potential_input_path = os.path.join(script_dir, "input", excel_filename)
                if os.path.exists(potential_input_path):
                    excel_path = potential_input_path
            
            # 如果在 input 目錄找不到，再檢查 output 目錄
            if not excel_path:
                output_dir = self.get_output_path()
                potential_output_path = os.path.join(output_dir, excel_filename)
                if os.path.exists(potential_output_path):
                    excel_path = potential_output_path
            
            if not excel_path:
                print(f"錯誤: 找不到 Excel 檔案: {excel_filename}")
                print(f"已搜尋的路徑:")
                if getattr(sys, 'frozen', False):
                    exe_dir = os.path.dirname(os.path.abspath(sys.argv[0]))
                    possible_paths = [
                        os.path.join(exe_dir, "..", "input", excel_filename),
                        os.path.join(exe_dir, "input", excel_filename),
                    ]
                    for path in possible_paths:
                        abs_path = os.path.abspath(path)
                        print(f"  - {abs_path} (存在: {os.path.exists(abs_path)})")
                return False
            
            print(f"✓ 找到 Excel 檔案: {excel_path}")
            
            # Word 檔案路徑
            output_dir = self.get_output_path()
            word_path = os.path.join(output_dir, word_filename)
            if not os.path.exists(word_path):
                print(f"錯誤: Word 檔案不存在: {word_path}")
                print(f"output 目錄: {output_dir}")
                print(f"目錄內容: {os.listdir(output_dir) if os.path.exists(output_dir) else '目錄不存在'}")
                return False
            
            print(f"✓ 找到 Word 檔案: {word_path}")
            
            # 讀取 Excel 檔案
            wb = openpyxl.load_workbook(excel_path, data_only=True)
            if sheet_name not in wb.sheetnames:
                print(f"錯誤: Excel 檔案中找不到工作表: {sheet_name}")
                available_sheets = ", ".join(wb.sheetnames)
                print(f"可用的工作表: {available_sheets}")
                return False
            
            ws = wb[sheet_name]
            
            # 開啟 Word 文件
            doc = Document(word_path)
            
            # 找到目標表格
            target_table = None
            if isinstance(table_identifier, int):
                # 使用精確索引
                if table_identifier >= len(doc.tables):
                    print(f"錯誤: 表格索引 {table_identifier} 超出範圍，文件共有 {len(doc.tables)} 個表格")
                    return False
                target_table = doc.tables[table_identifier]
                
            elif isinstance(table_identifier, str):
                # 使用文字搜尋
                table_index = self.find_table_by_content(doc, table_identifier)
                if table_index == -1:
                    print(f"錯誤: 找不到包含 '{table_identifier}' 的表格")
                    return False
                target_table = doc.tables[table_index]
                
            else:
                print(f"錯誤: 不支援的表格識別方式: {type(table_identifier)}")
                return False
            
            # 清除現有內容（如果需要）
            if clear_existing:
                for row in target_table.rows:
                    for cell in row.cells:
                        self._safe_set_text(cell, "")
                # 清除表格現有內容
                pass
            
            # 獲取要複製的數據範圍
            if cell_range:
                try:
                    excel_data = ws[cell_range]
                except Exception as e:
                    print(f"錯誤: 無效的儲存格範圍 '{cell_range}': {e}")
                    return False
            else:
                # 複製整個有效範圍
                max_row = ws.max_row
                max_col = ws.max_column
                if max_row == 1 and max_col == 1:
                    excel_data = ws['A1']
                else:
                    excel_data = ws[f"A1:{openpyxl.utils.get_column_letter(max_col)}{max_row}"]
            
            # 複製數據到 Word 表格
            start_row, start_col = start_position
            copied_cells = 0
            processed_cells = set()  # 記錄已處理的儲存格，避免重複處理
            
            if isinstance(excel_data, tuple):
                # 多行數據
                for row_idx, excel_row in enumerate(excel_data):
                    # 跳過 tech 表格的第四行，因為它已經被合併到第三行
                    if self._should_skip_row(row_idx):
                        continue
                        
                    # 計算 Word 中的實際行索引（考慮被跳過的行）
                    word_row_idx = start_row + row_idx
                    if self.current_sheet_name == 'tech' and row_idx > 3:
                        # 如果是 tech 表格且處理第五行以後的數據，需要向上偏移一行
                        word_row_idx -= 1
                    
                    if word_row_idx >= len(target_table.rows):
                        # 針對 tech 表格的特殊處理：由於第3、4行合併，理論上需要的行數會減少一行
                        if self.current_sheet_name == 'tech':
                            # tech 表格因為合併了一行，所以最後一行數據超出範圍是正常的
                            # 只有當明顯超出預期範圍時才警告
                            actual_needed_rows = len(target_table.rows) + 1  # tech 表格的實際需求行數
                            if word_row_idx >= actual_needed_rows:
                                print(f"警告: tech 表格數據過多，Excel第{row_idx+1}行無法放入 Word 表格")
                        else:
                            print(f"警告: Word 表格行數不足，已複製到第 {word_row_idx+1} 行")
                        break
                    
                    if isinstance(excel_row, tuple):
                        # 多列數據
                        for col_idx, excel_cell in enumerate(excel_row):
                            word_col_idx = start_col + col_idx
                            
                            if word_col_idx >= len(target_table.columns):
                                print(f"警告: Word 表格列數不足，已複製到第 {word_col_idx} 列")
                                break
                            
                            # 檢查是否已處理過此儲存格
                            cell_key = (word_row_idx, word_col_idx)
                            if cell_key in processed_cells:
                                continue
                            
                            # 檢查是否需要合併多行內容到同一個儲存格
                            merged = self._copy_cell_data_with_merge(excel_data, excel_cell, row_idx, col_idx, 
                                                                   target_table.cell(word_row_idx, word_col_idx))
                            
                            processed_cells.add(cell_key)
                            
                            # 內容模式合併：標記下一行
                            if merged and not self._should_merge_by_position(row_idx, col_idx):
                                next_cell_key = (word_row_idx + 1, word_col_idx)
                                processed_cells.add(next_cell_key)
                            
                            copied_cells += 1
                    else:
                        # 單列數據
                        if start_col < len(target_table.columns):
                            self._copy_cell_data(excel_row, target_table.cell(word_row_idx, start_col))
                            copied_cells += 1
            else:
                # 單個儲存格
                if (start_row < len(target_table.rows) and 
                    start_col < len(target_table.columns)):
                    self._copy_cell_data(excel_data, target_table.cell(start_row, start_col))
                    copied_cells += 1
            
            # 保存 Word 文件
            doc.save(word_path)
            # 操作成功完成
            pass
            return True
            
        except Exception as e:
            print(f"貼上 Excel 數據到 Word 表格時發生錯誤: {e}")
            import traceback
            traceback.print_exc()
            return False
        
        finally:
            try:
                if 'wb' in locals():
                    wb.close()
            except:
                pass
            # 清理當前工作表名稱
            if hasattr(self, 'current_sheet_name'):
                delattr(self, 'current_sheet_name')
    
    def list_word_tables(self, word_filename):
        """
        列出 Word 文件中所有表格的詳細資訊
        
        Args:
            word_filename: Word 檔案名稱（在 output 目錄中）
        
        Returns:
            list: 表格資訊列表
        """
        try:
            output_path = self.get_output_path()
            word_path = os.path.join(output_path, word_filename)
            
            if not os.path.exists(word_path):
                print(f"錯誤: Word 檔案不存在: {word_path}")
                return []
            
            doc = Document(word_path)
            table_info = []
            
            for i, table in enumerate(doc.tables):
                info = {
                    'index': i,
                    'rows': len(table.rows),
                    'columns': len(table.columns),
                    'first_cell_content': '',
                    'preview': ''
                }
                
                # 獲取第一個儲存格的內容
                if len(table.rows) > 0 and len(table.rows[0].cells) > 0:
                    first_cell = table.rows[0].cells[0].text.strip()
                    info['first_cell_content'] = first_cell[:50]
                
                # 獲取表格前兩行的內容預覽
                preview_lines = []
                for row_idx in range(min(2, len(table.rows))):
                    row_content = []
                    for cell_idx in range(min(3, len(table.rows[row_idx].cells))):
                        cell_text = table.rows[row_idx].cells[cell_idx].text.strip()[:15]
                        row_content.append(cell_text)
                    preview_lines.append(" | ".join(row_content))
                info['preview'] = " // ".join(preview_lines)
                
                table_info.append(info)
            
            return table_info
            
        except Exception as e:
            print(f"列出表格資訊時發生錯誤: {e}")
            return []
    
    def _copy_cell_data_with_merge(self, excel_data, excel_cell, row_idx, col_idx, word_cell):
        """
        複製儲存格數據，支援合併多行內容到同一個 Word 儲存格
        
        Args:
            excel_data: 完整的 Excel 數據範圍
            excel_cell: 當前處理的 Excel 儲存格
            row_idx: 當前行索引
            col_idx: 當前列索引
            word_cell: 目標 Word 儲存格
            
        Returns:
            bool: 是否進行了合併
        """
        # 檢查是否需要根據位置合併（例如：第三行和第四行）
        if self._should_merge_by_position(row_idx, col_idx):
            if (row_idx + 1 < len(excel_data) and isinstance(excel_data, tuple)):
                next_row = excel_data[row_idx + 1]
                if isinstance(next_row, tuple) and col_idx < len(next_row):
                    next_cell = next_row[col_idx]
                    
                    current_value = str(excel_cell.value) if excel_cell.value is not None else ""
                    next_value = str(next_cell.value) if next_cell.value is not None else ""
                    
                    # 合併第三行和第四行的內容
                    merged_content = f"{current_value}\n{next_value}"
                    self._safe_set_text(word_cell, merged_content)
                    
                    # 設定格式
                    for paragraph in word_cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Arial'
                            run.font.size = Pt(10)
                    return True
        
        # 檢查是否需要合併相鄰行的內容（例如：% of Q2 和 Revenue）
        current_value = str(excel_cell.value) if excel_cell.value is not None else ""
        
        # 如果當前行不是第一行，檢查是否需要與上一行合併
        if row_idx > 0 and isinstance(excel_data, tuple):
            prev_row = excel_data[row_idx - 1]
            if isinstance(prev_row, tuple) and col_idx < len(prev_row):
                prev_cell = prev_row[col_idx]
                prev_value = str(prev_cell.value) if prev_cell.value is not None else ""
                
                # 檢查是否是需要合併的模式（例如：% of Q2 和 Revenue）
                if self._should_merge_cells(prev_value, current_value):
                    # 合併兩行內容，用換行符分隔
                    merged_content = f"{prev_value}\n{current_value}"
                    self._safe_set_text(word_cell, merged_content)
                    
                    # 設定格式
                    for paragraph in word_cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Arial'
                            run.font.size = Pt(10)
                    return True
        
        # 檢查下一行是否需要與當前行合併
        if (row_idx < len(excel_data) - 1 and isinstance(excel_data, tuple) and 
            row_idx + 1 < len(excel_data)):
            next_row = excel_data[row_idx + 1]
            if isinstance(next_row, tuple) and col_idx < len(next_row):
                next_cell = next_row[col_idx]
                next_value = str(next_cell.value) if next_cell.value is not None else ""
                
                # 檢查是否是需要合併的模式
                if self._should_merge_cells(current_value, next_value):
                    # 合併兩行內容，用換行符分隔
                    merged_content = f"{current_value}\n{next_value}"
                    self._safe_set_text(word_cell, merged_content)
                    
                    # 設定格式
                    for paragraph in word_cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Arial'
                            run.font.size = Pt(10)
                    return True
        
        # 如果不需要合併，使用原有的複製方法
        self._copy_cell_data(excel_cell, word_cell)
        return False
    
    def _should_merge_cells(self, prev_value, current_value):
        """
        判斷是否應該合併兩個儲存格的內容
        
        Args:
            prev_value: 上一行的儲存格值
            current_value: 當前行的儲存格值
            
        Returns:
            bool: 是否應該合併
        """
        # 定義需要合併的模式
        merge_patterns = [
            ("% of Q2", "Revenue"),
            ("% of Q1", "Revenue"),
            ("% of Q3", "Revenue"),
            ("% of Q4", "Revenue"),
            ("% of", "Revenue"),  # 更寬泛的匹配
        ]
        
        prev_value_clean = prev_value.strip()
        current_value_clean = current_value.strip()
        
        for pattern_prev, pattern_current in merge_patterns:
            if (pattern_prev in prev_value_clean and 
                pattern_current in current_value_clean):
                return True
        
        return False
    
    def _should_merge_by_position(self, row_idx, col_idx):
        """
        判斷是否應該根據位置合併儲存格（只針對 tech 表格的第三行和第四行）
        
        Args:
            row_idx: 當前行索引
            col_idx: 當前列索引
            
        Returns:
            bool: 是否應該基於位置合併
        """
        # 只有 tech 表格才進行位置合併
        if not hasattr(self, 'current_sheet_name') or self.current_sheet_name != 'tech':
            return False
        
        # 只在第三行（索引2）時進行合併操作
        return row_idx == 2
    
    def _should_skip_row(self, row_idx):
        """
        判斷是否應該跳過某行（被合併的行）
        
        Args:
            row_idx: 當前行索引
            
        Returns:
            bool: 是否應該跳過
        """
        # 只有 tech 表格的第四行（索引3）需要跳過，因為它已經被合併到第三行
        if not hasattr(self, 'current_sheet_name') or self.current_sheet_name != 'tech':
            return False
        
        return row_idx == 3
    
    def _safe_set_text(self, word_cell, text_value):
        """
        安全地設置 Word 儲存格文本，處理特殊字符
        
        Args:
            word_cell: Word 儲存格物件
            text_value: 要設置的文本值
        """
        if text_value is None:
            text_value = ""
        
        # 將文本轉換為字符串
        text_str = str(text_value)
        
        # 直接設置文本，python-docx 會自動處理 XML 轉義
        word_cell.text = text_str
    
    def _copy_cell_data(self, excel_cell, word_cell):
        """
        複製單個儲存格的數據和格式，保留千分位分隔符和百分比符號
        
        Args:
            excel_cell: Excel 儲存格物件
            word_cell: Word 儲存格物件
        """
        if excel_cell.value is not None:
            # 嘗試獲取 Excel 儲存格的顯示文字（保留格式）
            try:
                # 檢查 Excel 儲存格是否有數字格式
                if hasattr(excel_cell, 'number_format') and excel_cell.number_format:
                    number_format = excel_cell.number_format
                    cell_value = excel_cell.value
                    
                    # 處理百分比格式
                    if '%' in number_format and isinstance(cell_value, (int, float)):
                        # Excel 中百分比值通常是小數，需要轉換為百分比顯示
                        # 但有些情況下可能已經是百分比形式（如 1 = 100%）
                        if abs(cell_value) <= 10:  # 假設小於等於10的數值是百分比形式
                            formatted_value = f"{cell_value * 100:.1f}%"
                        else:  # 大於10的數值可能已經是百分比形式
                            formatted_value = f"{cell_value:.1f}%"
                        self._safe_set_text(word_cell, formatted_value)
                        
                    # 處理千分位分隔符格式 - 金額取到整數
                    elif (',' in number_format or '#,##0' in number_format) and isinstance(cell_value, (int, float)):
                        # 所有金額都取到整數
                        formatted_value = f"{int(cell_value):,}"
                        self._safe_set_text(word_cell, formatted_value)
                        
                    # 處理貨幣格式 - 金額取到整數
                    elif ('$' in number_format or 'USD' in number_format or 'NT$' in number_format) and isinstance(cell_value, (int, float)):
                        currency_symbol = '$'
                        if 'NT$' in number_format:
                            currency_symbol = 'NT$'
                        
                        # 所有貨幣都取到整數
                        formatted_value = f"{currency_symbol}{int(cell_value):,}"
                        self._safe_set_text(word_cell, formatted_value)
                        
                    # 其他數字格式
                    elif isinstance(cell_value, (int, float)):
                        if isinstance(cell_value, float) and cell_value.is_integer():
                            self._safe_set_text(word_cell, str(int(cell_value)))
                        else:
                            self._safe_set_text(word_cell, str(cell_value))
                    else:
                        self._safe_set_text(word_cell, str(cell_value))
                else:
                    # 沒有特殊格式的情況
                    cell_value = excel_cell.value
                    if isinstance(cell_value, (int, float)):
                        if isinstance(cell_value, float) and cell_value.is_integer():
                            self._safe_set_text(word_cell, str(int(cell_value)))
                        else:
                            self._safe_set_text(word_cell, str(cell_value))
                    else:
                        self._safe_set_text(word_cell, str(cell_value))
                        
            except Exception as e:
                # 如果格式化失敗，使用原始值
                # 格式化警告，使用原始值
                pass
                self._safe_set_text(word_cell, str(excel_cell.value))
            
            # 設定 Word 格式
            for paragraph in word_cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Arial'  # 設定為 Arial 字體
                    run.font.size = Pt(10)
                    # 不使用粗體格式
    
    def batch_paste_excel_to_word(self, paste_configs):
        """
        批量貼上多個 Excel 範圍到 Word 表格
        
        Args:
            paste_configs: 貼上配置列表，每個配置包含:
                {
                    'excel_filename': str,
                    'sheet_name': str,
                    'word_filename': str,
                    'table_identifier': int or str,
                    'cell_range': str (optional),
                    'start_position': tuple (optional),
                    'clear_existing': bool (optional)
                }
        
        Returns:
            dict: 每個配置的執行結果
        """
        results = {}
        
        for i, config in enumerate(paste_configs):
            config_name = f"配置_{i+1}"
            
            try:
                result = self.paste_excel_to_word_table(
                    excel_filename=config.get('excel_filename'),
                    sheet_name=config.get('sheet_name'),
                    word_filename=config.get('word_filename'),
                    table_identifier=config.get('table_identifier'),
                    cell_range=config.get('cell_range'),
                    start_position=config.get('start_position', (0, 0)),
                    clear_existing=config.get('clear_existing', False)
                )
                results[config_name] = result
                
            except Exception as e:
                print(f"{config_name} 執行失敗: {e}")
                results[config_name] = False
        
        return results


# 便利函數
def create_excel_table2word():
    """創建 Excel 表格到 Word 的整合實例"""
    return ExcelTable2Word()


def quick_paste_excel_to_word(excel_filename, sheet_name, word_filename, 
                             table_identifier, cell_range=None, 
                             start_position=(0, 0)):
    """
    快速貼上 Excel 表格數據到 Word 表格的便利函數
    
    Args:
        excel_filename: Excel 檔案名稱
        sheet_name: Excel 工作表名稱
        word_filename: Word 檔案名稱
        table_identifier: 表格識別（索引或文字）
        cell_range: Excel 儲存格範圍
        start_position: Word 表格開始位置
    
    Returns:
        bool: 操作是否成功
    """
    integration = ExcelTable2Word()
    return integration.paste_excel_to_word_table(
        excel_filename, sheet_name, word_filename, 
        table_identifier, cell_range, start_position
    )



