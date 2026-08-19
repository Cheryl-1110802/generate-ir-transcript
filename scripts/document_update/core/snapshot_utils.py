"""
Quarter snapshot utilities for stale-data detection.

A snapshot records the MD5 hash of each Excel sheet, keyed by
"{year}_{quarter}".  The flow is:

  1. User calls /update-config  → take_snapshot() stores hashes of the
     CURRENT (last-quarter) data before the new quarter data arrives.

  2. User fills in new data and calls /run/script1
     → get_stale_sheets() compares current hashes against the snapshot.
       Sheets whose hash hasn't changed are "stale" (still last-quarter data).

  3. build_context() calls mark_stale_values() on each stale sheet's
     parsed dict, prepending an invisible PUA sentinel to every string.

  4. After generate_report() writes the Word file, color_stale_paragraphs()
     opens it with python-docx, colors every paragraph that contains the
     sentinel red, removes the sentinel, and saves.

First-run behaviour (no snapshot file or no entry for this quarter):
  → get_stale_sheets() returns an empty set → all text stays black.
"""

import hashlib
import json
import os
import pandas as pd
from datetime import datetime

# Unicode Private-Use-Area char – invisible, won't appear in normal text
STALE_MARKER: str = ""
SNAPSHOT_FILENAME: str = "quarter_history.json"
STALE_COLOR: str = "FF0000"   # red


# ── hashing ──────────────────────────────────────────────────────────────────

def compute_sheet_hash(df: pd.DataFrame) -> str:
    """Stable MD5 hash of a DataFrame (index + values)."""
    try:
        raw = pd.util.hash_pandas_object(df, index=True).values.tobytes()
        return hashlib.md5(raw).hexdigest()
    except Exception:
        return hashlib.md5(df.to_csv().encode("utf-8", errors="replace")).hexdigest()


# ── snapshot I/O ──────────────────────────────────────────────────────────────

def take_snapshot(excel_path: str, sheet_names: list,
                  quarter: str, year: str, snapshot_path: str) -> dict:
    """
    Hash every sheet in *excel_path* and persist under ``"{year}_{quarter}"``
    in *snapshot_path*.  Returns the ``{sheet: hash}`` dict stored.

    Called from /update-config when the user starts a new quarter.
    """
    try:
        xls = pd.ExcelFile(excel_path)
    except Exception as e:
        raise IOError(f"無法讀取 Excel 快照來源: {e}")

    hashes: dict = {}
    for sheet in sheet_names:
        try:
            df = xls.parse(sheet, index_col=0, header=0)
            hashes[sheet] = compute_sheet_hash(df)
        except Exception:
            hashes[sheet] = None   # sheet absent / unreadable

    key = f"{year}_{quarter}"

    data: dict = {}
    if os.path.exists(snapshot_path):
        try:
            with open(snapshot_path, "r", encoding="utf-8") as f:
                data = json.load(f)
        except Exception:
            data = {}

    data[key] = {
        "quarter": quarter,
        "year": str(year),
        "snapshot_time": datetime.now().isoformat(),
        "sheet_hashes": hashes,
    }

    dir_ = os.path.dirname(snapshot_path)
    if dir_:
        os.makedirs(dir_, exist_ok=True)
    with open(snapshot_path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)

    return hashes


def get_stale_sheets(excel_path: str, sheet_names: list,
                     quarter: str, year: str, snapshot_path: str) -> set:
    """
    Return the set of sheet names whose hash is *unchanged* since the
    snapshot (same hash ⟹ data not updated ⟹ stale ⟹ will be red).

    Returns an empty set if no snapshot exists (first run → all black).
    """
    if not os.path.exists(snapshot_path):
        return set()

    key = f"{year}_{quarter}"
    try:
        with open(snapshot_path, "r", encoding="utf-8") as f:
            data = json.load(f)
    except Exception:
        return set()

    if key not in data:
        return set()

    snapshot_hashes: dict = data[key].get("sheet_hashes", {})

    try:
        xls = pd.ExcelFile(excel_path)
    except Exception:
        return set()

    stale: set = set()
    for sheet in sheet_names:
        prev_hash = snapshot_hashes.get(sheet)
        if prev_hash is None:
            continue
        try:
            df = xls.parse(sheet, index_col=0, header=0)
            if compute_sheet_hash(df) == prev_hash:
                stale.add(sheet)
        except Exception:
            pass

    return stale


# ── stale-value marking ───────────────────────────────────────────────────────

def mark_stale_values(data, marker: str = STALE_MARKER):
    """
    Recursively prepend *marker* to every string / number value so that
    color_stale_paragraphs() can locate runs from stale sheets in the
    rendered Word document.
    """
    if isinstance(data, dict):
        return {k: mark_stale_values(v, marker) for k, v in data.items()}
    elif isinstance(data, list):
        return [mark_stale_values(item, marker) for item in data]
    elif isinstance(data, str):
        return marker + data
    elif isinstance(data, (int, float)):
        return marker + str(data)   # stringify so marker can be prepended
    else:
        return data                 # None, bool, etc. – leave as-is


# ── post-render coloring ──────────────────────────────────────────────────────

def color_stale_paragraphs(docx_path: str,
                            marker: str = STALE_MARKER,
                            hex_color: str = STALE_COLOR) -> None:
    """
    Open the rendered Word document, color RED every paragraph (body or
    table-cell) that contains *marker*, strip the markers, and save.

    A paragraph is colored entirely red when any run inside it carries the
    sentinel – even the surrounding template text is red, since the whole
    sentence is based on possibly-stale numbers.
    """
    from docx import Document
    from docx.shared import RGBColor

    r = int(hex_color[0:2], 16)
    g = int(hex_color[2:4], 16)
    b = int(hex_color[4:6], 16)
    color = RGBColor(r, g, b)

    doc = Document(docx_path)

    def _process_para(para):
        if any(marker in run.text for run in para.runs):
            for run in para.runs:
                run.font.color.rgb = color
                run.text = run.text.replace(marker, "")

    for para in doc.paragraphs:
        _process_para(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _process_para(para)

    doc.save(docx_path)
